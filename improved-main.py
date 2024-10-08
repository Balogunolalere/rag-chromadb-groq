# Imports
import streamlit as st
import groq
import os
import json
import time
from typing import List, Dict, Tuple, Any
from dotenv import load_dotenv
from io import BytesIO
import logging
from functools import lru_cache

import spacy
from langdetect import detect
from textstat import flesch_reading_ease

import chromadb
from chromadb.utils import embedding_functions

from PyPDF2 import PdfReader
import docx


# Configuration
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

load_dotenv()

REQUIRED_ENV_VARS = ["GROQ_API_KEY"]
PERSIST_DIRECTORY = "chromadb_data"
COLLECTION_NAME = "document_collection"
EMBEDDING_MODEL = 'all-MiniLM-L6-v2'
LLM_MODEL = "llama-3.1-70b-versatile"

# Initialization
def load_spacy_model():
    try:
        return spacy.load("en_core_web_lg")
    except OSError:
        logger.info("Downloading spaCy model. This may take a while...")
        spacy.cli.download("en_core_web_lg")
        return spacy.load("en_core_web_lg")

nlp = load_spacy_model()

def check_env_vars():
    missing_vars = [var for var in REQUIRED_ENV_VARS if not os.getenv(var)]
    if missing_vars:
        raise EnvironmentError(f"Missing required environment variables: {', '.join(missing_vars)}")

check_env_vars()

client = groq.Groq(api_key=os.getenv("GROQ_API_KEY"))
chroma_client = chromadb.PersistentClient(path=PERSIST_DIRECTORY)
collection = chroma_client.get_or_create_collection(
    name=COLLECTION_NAME,
    embedding_function=embedding_functions.SentenceTransformerEmbeddingFunction(model_name=EMBEDDING_MODEL)
)

# Caching
@lru_cache(maxsize=100)
def cached_api_call(messages: Tuple[Tuple[str, str]], max_tokens: int) -> str:
    messages_list = [{"role": role, "content": content} for role, content in messages]
    response = client.chat.completions.create(
        model=LLM_MODEL,
        messages=messages_list,
        max_tokens=max_tokens,
        temperature=0.2
    )
    return response.choices[0].message.content

# Updated HybridChunker Class
class HybridChunker:
    def __init__(self, max_chunk_size: int = 1000, overlap: int = 200, similarity_threshold: float = 0.5,
                 use_semantic: bool = True, use_hierarchical: bool = True, adaptive_chunking: bool = True):
        self.max_chunk_size = max_chunk_size
        self.overlap = overlap
        self.similarity_threshold = similarity_threshold
        self.use_semantic = use_semantic
        self.use_hierarchical = use_hierarchical
        self.adaptive_chunking = adaptive_chunking
        self.nlp = spacy.load("en_core_web_lg")

    def chunk_text(self, text: str) -> List[str]:
        doc = self.nlp(text)
        sentences = list(doc.sents)
        chunks = []
        current_chunk = []
        current_size = 0

        for i, sentence in enumerate(sentences):
            sentence_size = len(sentence.text.split())
            
            # Adaptive Chunk Sizing
            if self.adaptive_chunking:
                chunk_size_limit = self._calculate_adaptive_chunk_size(sentence.text)
            else:
                chunk_size_limit = self.max_chunk_size

            if current_size + sentence_size > chunk_size_limit:
                if current_chunk:
                    chunks.append(self._process_chunk(current_chunk))
                    overlap_size = sum(len(s.text.split()) for s in current_chunk[-2:])
                    current_chunk = current_chunk[-2:] if len(current_chunk) > 2 else []
                    current_size = overlap_size
                else:
                    current_size = 0

            current_chunk.append(sentence)
            current_size += sentence_size

            # Check for semantic similarity
            if self.use_semantic and i < len(sentences) - 1:
                next_sentence = sentences[i + 1]
                similarity = sentence.similarity(next_sentence)
                if similarity < self.similarity_threshold and current_size >= chunk_size_limit // 2:
                    chunks.append(self._process_chunk(current_chunk))
                    overlap_size = sum(len(s.text.split()) for s in current_chunk[-2:])
                    current_chunk = current_chunk[-2:] if len(current_chunk) > 2 else []
                    current_size = overlap_size

        if current_chunk:
            chunks.append(self._process_chunk(current_chunk))

        # Hierarchical chunking if enabled
        if self.use_hierarchical:
            return self._hierarchical_chunking(chunks)
        else:
            return chunks

    def _process_chunk(self, chunk: List[spacy.tokens.Span]) -> str:
        """Joins sentences to form a chunk, preserving paragraph structure."""
        return " ".join([s.text for s in chunk])

    def _calculate_adaptive_chunk_size(self, text: str) -> int:
        """Uses readability score or sentence complexity to adjust chunk size."""
        score = flesch_reading_ease(text)
        if score < 60:  # More complex text, reduce chunk size
            return int(self.max_chunk_size * 0.8)
        else:  # Simpler text, allow larger chunks
            return self.max_chunk_size

    def _hierarchical_chunking(self, chunks: List[str]) -> List[str]:
        """Implements hierarchical chunking with larger overlapping chunks."""
        hierarchical_chunks = []
        for i in range(len(chunks)):
            # Create larger overlapping chunks
            if i + 2 < len(chunks):
                larger_chunk = " ".join([chunks[i], chunks[i + 1], chunks[i + 2]])
                hierarchical_chunks.append(larger_chunk)
        return hierarchical_chunks + chunks

# Document Processing
def process_pdf(file: BytesIO) -> Tuple[str, Dict[str, Any]]:
    text = ""
    metadata = {"pages": []}

    try:
        pdf = PdfReader(file)
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            text += page_text + "\n"
            metadata["pages"].append({"number": i + 1, "text": page_text})

        metadata["num_pages"] = len(pdf.pages)
        if pdf.metadata:
            for key, value in pdf.metadata.items():
                if key != "pages":
                    metadata[key] = value

    except Exception as e:
        logger.error(f"Error processing PDF: {str(e)}")
        raise

    return text, metadata

def process_docx(file: BytesIO) -> Tuple[str, Dict[str, Any]]:
    text = ""
    metadata = {"paragraphs": []}

    try:
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
            metadata["paragraphs"].append(para.text)

        metadata["num_paragraphs"] = len(doc.paragraphs)
        metadata["core_properties"] = {
            prop: getattr(doc.core_properties, prop)
            for prop in dir(doc.core_properties)
            if not prop.startswith("_") and prop != "custom_properties"
        }

    except Exception as e:
        logger.error(f"Error processing DOCX: {str(e)}")
        raise

    return text, metadata

def process_json(file: BytesIO) -> Tuple[str, Dict[str, Any]]:
    try:
        content = json.load(file)
        text = json.dumps(content, indent=2)
        metadata = {
            "top_level_keys": list(content.keys()) if isinstance(content, dict) else [],
            "size": len(text)
        }
        return text, metadata
    except json.JSONDecodeError as e:
        logger.error(f"Error processing JSON: {str(e)}")
        raise

def safe_file_read(file, fallback_encoding='latin1'):
    try:
        return file.read().decode('utf-8')
    except UnicodeDecodeError:
        file.seek(0)
        return file.read().decode(fallback_encoding)

def clean_text(text: str) -> str:
    text = " ".join(text.split())
    return ''.join(char for char in text if char.isprintable() or char.isspace())
# Storage Management
def clean_storage():
    global chroma_client, collection

    chroma_client.delete_collection(name=COLLECTION_NAME)
    chroma_client = chromadb.PersistentClient(path=PERSIST_DIRECTORY)
    collection = chroma_client.get_or_create_collection(
        name=COLLECTION_NAME,
        embedding_function=embedding_functions.SentenceTransformerEmbeddingFunction(model_name=EMBEDDING_MODEL)
    )

    st.session_state.document_metadata = {}

    if os.path.exists("history.json"):
        os.remove("history.json")

    st.success("Storage cleaned successfully!")

# Document Upload
def upload_document():
    uploaded_file = st.file_uploader("Upload a PDF, TXT, JSON, or DOCX file", type=["pdf", "txt", "json", "docx"])

    if uploaded_file is not None:
        if uploaded_file.name in st.session_state.document_metadata:
            st.warning(f"Document '{uploaded_file.name}' has already been processed.")
            return

        try:
            if uploaded_file.type == "application/pdf":
                text, metadata = process_pdf(BytesIO(uploaded_file.getvalue()))
            elif uploaded_file.type == "text/plain":
                text = safe_file_read(uploaded_file)
                metadata = {"type": "txt"}
            elif uploaded_file.type == "application/json":
                text, metadata = process_json(BytesIO(uploaded_file.getvalue()))
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text, metadata = process_docx(BytesIO(uploaded_file.getvalue()))
            else:
                st.error(f"Unsupported file type: {uploaded_file.type}")
                return

            text = clean_text(text)

            try:
                language = detect(text)
            except:
                language = "unknown"

            metadata.update({
                "language": language,
                "file_name": uploaded_file.name,
                "file_size": uploaded_file.size,
                "file_type": uploaded_file.type
            })

            # Using the updated HybridChunker
            chunker = HybridChunker()
            passages = chunker.chunk_text(text)

            ids = [f"{uploaded_file.name}_{i}" for i in range(len(passages))]

            collection.add(
                ids=ids,
                documents=passages,
                metadatas=[{"chunk_id": i, "source": uploaded_file.name} for i in range(len(passages))]
            )

            st.session_state.document_metadata[uploaded_file.name] = metadata

            st.success(f"Document '{uploaded_file.name}' uploaded and processed successfully!")
            st.write("Document Details:")
            st.json(metadata)

        except Exception as e:
            st.error(f"An error occurred while processing the file: {str(e)}")
            logger.error(f"Error in upload_document: {str(e)}", exc_info=True)

# Query Processing
def retrieve_relevant_passages(query: str) -> List[str]:
    results = collection.query(
        query_texts=[query],
        n_results=5
    )
    return results['documents'][0]

def make_api_call(messages: List[Dict[str, str]], max_tokens: int, is_final_answer: bool = False) -> str:
    messages_tuple = tuple((msg["role"], msg["content"]) for msg in messages)

    for attempt in range(3):
        try:
            return cached_api_call(messages_tuple, max_tokens)
        except Exception as e:
            logger.error(f"API call attempt {attempt + 1} failed: {str(e)}")
            if attempt == 2:
                error_message = f"Failed to generate {'final answer' if is_final_answer else 'step'} after 3 attempts. Error: {str(e)}"
                logger.error(error_message)
                return json.dumps({"title": "Error", "content": error_message, "next_action": "final_answer" if not is_final_answer else None})
            time.sleep(1)

def generate_response(prompt: str, retrieved_context: str):
    messages = [
        {"role": "system", "content": "You are an expert AI assistant. Use the provided context to generate your answer. Follow these steps: 1) Analyze the question. 2) Identify key information from the context. 3) Reason about the answer. 4) Formulate a clear response. Your response for each step MUST be in valid JSON format with 'step_number', 'title', 'content', and 'next_action' fields. The 'content' field MUST be a string. The 'next_action' field should be 'continue' for steps 1-3 and 'final_answer' for step 4."},
        {"role": "user", "content": f"Context: {retrieved_context}\n\nQuestion: {prompt}\n\nBegin with step 1: Analyze the question. Respond in strict JSON format as instructed."}
    ]

    steps = []
    total_thinking_time = 0

    for step_number in range(1, 5):
        start_time = time.time()
        step_response = make_api_call(messages, 300)
        end_time = time.time()
        thinking_time = end_time - start_time
        total_thinking_time += thinking_time

        try:
            step_data = json.loads(step_response)
            if not isinstance(step_data.get('content'), str):
                raise ValueError("'content' field must be a string")
            if 'next_action' not in step_data or step_data['next_action'] not in ['continue', 'final_answer']:
                raise ValueError("Invalid or missing 'next_action' field")
            if int(step_data.get('step_number', 0)) != step_number:
                raise ValueError(f"Expected step {step_number}, got step {step_data.get('step_number')}")
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse JSON: {step_response}")
            step_data = {"step_number": step_number, "title": "Error", "content": f"Failed to parse response: {str(e)}", "next_action": "final_answer"}
        except ValueError as e:
            logger.error(f"Invalid JSON structure: {str(e)}")
            step_data = {"step_number": step_number, "title": "Error", "content": str(e), "next_action": "final_answer"}

        steps.append((f"Step {step_number}: {step_data['title']}", step_data['content'], thinking_time))

        if step_number < 4:
            messages.append({"role": "assistant", "content": json.dumps(step_data)})
            next_step_prompts = [
                "Now, proceed with step 2: Identify key information from the context.",
                "Next, move to step 3: Reason about the answer.",
                "Finally, complete step 4: Formulate a clear response."
            ]
            messages.append({"role": "user", "content": f"{next_step_prompts[step_number - 1]} Respond in strict JSON format as instructed."})
        else:
            return steps, total_thinking_time, step_data['content']

    return steps, total_thinking_time, "Error: Failed to generate a final answer."

# History Management
def save_to_history(query: str, response: str):
    history_entry = {"query": query, "response": response, "timestamp": time.time()}
    try:
        with open("history.json", "r+") as f:
            content = f.read()
            history = json.loads(content) if content else []
            history.append(history_entry)
            f.seek(0)
            json.dump(history, f, indent=2)
            f.truncate()
    except FileNotFoundError:
        with open("history.json", "w") as f:
            json.dump([history_entry], f, indent=2)

def show_history():
    try:
        with open("history.json", "r") as f:
            history = json.load(f)

        history.sort(key=lambda x: x['timestamp'], reverse=True)

        for entry in history[:10]:
            st.write(f"**Query**: {entry['query']}")
            st.write(f"**Response**: {entry['response']}")
            st.write(f"**Time**: {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(entry['timestamp']))}")
            if 'file_type' in entry:
                st.write(f"**File Type**: {entry['file_type']}")
            st.write("---")
    except FileNotFoundError:
        st.write("No history available.")

# Document Management
def view_uploaded_documents():
    if not st.session_state.document_metadata:
        st.write("No documents have been uploaded yet.")
    else:
        for filename, metadata in st.session_state.document_metadata.items():
            st.write(f"**Filename**: {filename}")
            st.write(f"**File Type**: {metadata.get('file_type', 'Unknown')}")
            st.write(f"**File Size**: {metadata.get('file_size', 'Unknown')} bytes")
            st.write(f"**Language**: {metadata.get('language', 'Unknown')}")
            if 'num_pages' in metadata:
                st.write(f"**Number of Pages**: {metadata['num_pages']}")
            elif 'num_paragraphs' in metadata:
                st.write(f"**Number of Paragraphs**: {metadata['num_paragraphs']}")
            st.write("---")

# Main Application
def main():
    st.set_page_config(page_title="RAG with ChromaDB", page_icon="🧠", layout="wide")

    st.title("RAG with ChromaDB and Groq")

    st.markdown("""
    <style>
    .big-font {
        font-size:20px !important;
        font-weight: bold;
    }
    .stButton>button {
        background-color: #4CAF50;
        color: white;
        padding: 10px 24px;
        border: none;
        border-radius: 4px;
        cursor: pointer;
    }
    .stButton>button:hover {
        background-color: #45a049;
    }
    </style>
    """, unsafe_allow_html=True)

    st.markdown('<p class="big-font">Upload PDF, TXT, JSON, or DOCX files, create a searchable knowledge base, and query documents using retrieval-augmented generation.</p>', unsafe_allow_html=True)

    col1, col2, col3 = st.columns([2, 1, 1])

    with col1:
        upload_document()

    with col2:
        if st.button("Clean Storage"):
            clean_storage()

    with col3:
        if st.button("View Uploaded Documents"):
            view_uploaded_documents()

    user_query = st.text_input("Enter your query:", placeholder="e.g., What is the main idea of the document?")
    submit_button = st.button("Submit Query")

    if submit_button and user_query:
        retrieved_passages = retrieve_relevant_passages(user_query)
        context = " ".join(retrieved_passages)

        with st.spinner("Generating response..."):
            try:
                steps, total_thinking_time, final_answer = generate_response(user_query, context)

                for title, content, thinking_time in steps:
                    with st.expander(title, expanded=True):
                        st.markdown(content.replace('\n', '<br>'), unsafe_allow_html=True)
                        st.write(f"Thinking time: {thinking_time:.2f} seconds")

                st.markdown("### Final Answer")
                st.markdown(final_answer.replace('\n', '<br>'), unsafe_allow_html=True)

                save_to_history(user_query, final_answer)
                st.success("Response generated successfully!")
                st.write(f"Total thinking time: {total_thinking_time:.2f} seconds")
            except Exception as e:
                st.error(f"An error occurred while generating the response: {str(e)}")
                logger.error(f"Error in generate_response: {str(e)}", exc_info=True)

    if st.button("Show History"):
        show_history()

if __name__ == "__main__":
    if 'document_metadata' not in st.session_state:
        st.session_state.document_metadata = {}
    main()
